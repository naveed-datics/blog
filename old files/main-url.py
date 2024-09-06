import streamlit as st
import requests
from bs4 import BeautifulSoup
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain_community.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
from langchain_ollama.llms import OllamaLLM
from docx import Document  # Import for creating a Word document

# from transformers import AutoTokenizer, AutoModelForCausalLM

load_dotenv()
os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

def get_website_text(urls):
    texts = []
    for url in urls:
        response = requests.get(url)
        soup = BeautifulSoup(response.content, 'html.parser')
        # Extract text from paragraphs and headings
        paragraphs = soup.find_all(['ul','li','ol','p','a','table','h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
        text = ' '.join([para.get_text() for para in paragraphs])
        texts.append(text)
    return ' '.join(texts)  # Combine texts from all URLs

def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=250)  # Adjusted sizes
    chunks = text_splitter.split_text(text)
    return chunks

def get_vector_store(text_chunks):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
    vector_store.save_local("faiss_index")

def get_conversational_chain():
    prompt_template = """
    Answer the question as detailed as possible from the provided context, make sure to provide all the details.
    Context:\n {context}?\n
    Question: focus keyword = {question}   
    
    Answer:
    """
    model = ChatGoogleGenerativeAI(model="gemini-1.5-flash-001", temperature=0.5,max_tokens=6500)
    # model = OllamaLLM(model="llama3",temperature=0.3, max_tokens=6500)
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])

    chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)
    return chain

def user_input(user_question):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/text-embedding-004")
    new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
    docs = new_db.similarity_search(user_question)
    chain = get_conversational_chain()
    response = chain({"input_documents": docs, "question": user_question}, return_only_outputs=True)
    print(response)
    st.write("Reply: ", response["output_text"])

def export_to_docx(text, filename="exported_content.docx"):
    document = Document()
    document.add_heading('Extracted Website Content', 0)
    document.add_paragraph(text)
    document.save(filename)
    return filename

def main():
    st.set_page_config("Chat Web Pages")
    st.header("Chat with Web Page Content")

    # Text area for adding a prompt/question
    user_question = st.text_area("Prompt", height=70)

    # Add a submit button to trigger the response generation
    if st.button("Generate Article"):
        if user_question:
            user_input(user_question)
        else:
            st.warning("Please enter a question before submitting.")

    with st.sidebar:
        st.title("Website URLs:")
        urls_input = st.text_area("Enter Website URLs (one per line)")
        urls = [url.strip() for url in urls_input.split('\n') if url.strip()]
        if st.button("Submit & Process"):
            with st.spinner("Processing..."):
                raw_text = get_website_text(urls)
                text_chunks = get_text_chunks(raw_text)
                get_vector_store(text_chunks)
                st.success("Processing Complete")

                # Export the website content to a Word document
                docx_file = export_to_docx(raw_text)
                st.success(f"Content exported to {docx_file}")
                with open(docx_file, "rb") as file:
                    st.download_button(label="Download Word Document", data=file, file_name=docx_file)

if __name__ == "__main__":
    if "logged_in" not in st.session_state:
        # Enable Password
        st.session_state.logged_in = True

    if not st.session_state.logged_in:

        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            col4, col5, col6 = st.columns([1, 1, 1])
            with col5:
                st.subheader("Login")
            email = st.text_input("Enter your email:")
            password = st.text_input("Enter your password:", type="password")
            col7, col8, col9 = st.columns([4, 4, 4])
            with col8:
                if st.button("Login"):
                    if email == 'naveed' and password == '123':
                        st.session_state.logged_in = True
                        st.rerun()
                    else:
                        st.toast('Please write the email and password yourself and do not use autofill')

    else:
        main()
  
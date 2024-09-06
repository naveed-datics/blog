import streamlit as st
import requests
from bs4 import BeautifulSoup
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain_community.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
from langchain_ollama.llms import OllamaLLM
from docx import Document  # Import for creating a Word document
from langchain_openai import ChatOpenAI
import json
from langchain.document_loaders import WebBaseLoader

load_dotenv()
os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

def sendWordPrssApi(title, text):
    url = "https://howdidcelebdie.com/wp-json/wp/v2/posts/"
    payload = json.dumps({
        "title": title,
        "content": text,
        "status": "draft",
        "slug": title,
        "meta": {
            "rank_math_title": title,
            "rank_math_focus_keyword": title,
        }
    })
    headers = {
        'Authorization': 'Basic YWRtaW46YWRtaW5Ad29yazEyMw==',
        'Content-Type': 'application/json'
    }
    response = requests.request("POST", url, headers=headers, data=payload)
    return response.status_code, response.text  # Return status and response text

def sendWordPrssApiTraslated(title, text):
    url = "https://howdidcelebdie.com/de/wp-json/wp/v2/posts"
    payload = json.dumps({
        "title": title,
        "content": text,
        "status": "draft",
        "slug": title,
        "meta": {
            "rank_math_title": title,
            "rank_math_focus_keyword": title,
        }
    })
    headers = {
        'Authorization': 'Basic YWRtaW46YWRtaW5Ad29yazEyMw==',
        'Content-Type': 'application/json'
    }
    response = requests.request("POST", url, headers=headers, data=payload)
    return response.status_code, response.text  # Return status and response text


def get_website_text(urls):
    loader = WebBaseLoader(urls)
    documents = loader.load()
    texts = [doc.page_content for doc in documents]
    return ' '.join(texts)  # Combine texts from all URLs

def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=250)  # Adjusted sizes
    chunks = text_splitter.split_text(text)
    return chunks

def get_vector_store(text_chunks):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
    vector_store.save_local("faiss_index")

def get_conversational_chain():
    prompt_template = """
    Answer the question as detailed as possible from the provided context, make sure to provide all the details.
    Context:\n {context}?\n
    Question: focus keyword = {question}   

   Create a comprehensive information about 3000 words on this {question} also includes every element of his bigraphy {question} from data given to you make me comprehensive article finsh cotent gap and merge all data in 1 article SEO-optimized article in format for WordPress

    Write Content in this order:
    1.	Start text with: Write 50 words summery and make usre use this {question} and clearly state the answer .'
    2.	Give me answers in html table with two colums 
        Religion: 
        Profession:        
        Date of birth: 
        Zodiac sign:  
        Nationality: 
    3.	Write as Frenklen an expert with 15 years of experience and have the authority to write this article in the introduction including expert opinions and high engagement tips.
    4.	{question} and Faith (600 words). 
    5.	{question} Parents (250 wordss).  heading redner in h2
    6.	Life Paretner religion
    7.	Education and Career Beginnings
    8.	Rise to Career
    9.	Collaborations
    10.	Personal Life and Relationships heading redner in h2 and avoid tag display in text
    11.	Net Worth and Career Success
    12. Influence and Legacy
    13.	Include if news about Attending Religious Events.   
    17. Related Quries and their answers and Quries render in H3
    18. Longtail keywrods and their answers and keywrods render in H3  and avoid tag display in text
    19. Conclusion heading redner in h2 and avoid tag display in text
    20. FAQs heading redner in h2 (10 with answers and Questions render in H3 and avoid tag display in text)
    21.	If you’re interested in learning more about religion, feel free to visit my website: whatreligionisinfo.com.
    

    SEO Optimization:
    1. Incorporate LSI and NLP-friendly keywords throughout for the focus keyword.
    2. Include a link to the Wikipedia page for the focus keyword.

    Additional Instructions:
    0.  Format the output in Markdown for WordPress.
    1.	Omit a title or introduction heading.
    3.	Ensure accuracy of all facts, figures, and representations.
    4.	Balance informative content with readability and engagement.
    5.	If mentioning specific individuals, avoid implying recognition or identification based on appearance.
    6.  Find focus keyword wikipida page and link in the text 
    9.  Do not include a title or introduction heading.
    11. Add Related Queries and Long tail Keywords and their answers

 
    Answer:
    """
    model = ChatGoogleGenerativeAI(model="gemini-1.5-flash-001", temperature=0.5, max_tokens=6500)
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)
    return chain


def translate_text(title, input_text, target_language):
    try:
        llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash-001", temperature=0.5, max_tokens=6500)
        
        # Translate the title
        title_messages = [
            ("system", f"Translate the following text string to {target_language}:"),
            ("human", title),
        ]
        title_response = llm.invoke(title_messages)
        translated_title = title_response.content
        
        # Translate the input_text
        text_messages = [
            ("system", f"Translate the following text to {target_language}:"),
            ("human", input_text),
        ]
        text_response = llm.invoke(text_messages)
        translated_text = text_response.content
        
        return translated_title, translated_text

    except Exception as e:
        print(f"Error during translation: {e}")
        return "Translation error occurred.", "Translation error occurred."


def user_input(user_question):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/text-embedding-004")
    new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
    docs = new_db.similarity_search(user_question)
    chain = get_conversational_chain()
    response = chain({"input_documents": docs, "question": user_question}, return_only_outputs=True)
    title = user_question
    text = response["output_text"]
    print(response)
    return title, text

def export_to_docx(text, filename="exported_content.docx"):
    document = Document()
    document.add_heading('Extracted Website Content', 0)
    document.add_paragraph(text)
    document.save(filename)
    return filename

def main():
    st.set_page_config("Chat Web Pages")
    st.header("Chat with Web Page Content")

    # Text area for adding a prompt/question
    user_question = st.text_area("Prompt", height=70)

    # Add a submit button to trigger the response generation
    if st.button("Generate Article"):
        if user_question:
            title, text = user_input(user_question)
            # translateTitle, translateText = translate_text(title,text, target_language="German")  # 'German' for German translation

            # st.session_state['title'] = title
            # st.session_state['text'] = text
            # st.session_state['t_title'] = translateTitle
            # st.session_state['t_text'] = translateText
            
            st.write(text)
            # st.write("---------------------------------------------------------------------------------")
            # st.markdown(translateTitle, unsafe_allow_html=True)
            # st.markdown(translateText, unsafe_allow_html=True)
            # st.success("Article generated successfully!")
        else:
            st.warning("Please enter a question before submitting.")

    # if 'title' in st.session_state and 'text' in st.session_state:
    #     if st.button("Post to WordPress"):
    #         status_code, response_text = sendWordPrssApi(st.session_state['title'], st.session_state['text'])
    #         if status_code == 201:
    #             st.success("English Article posted to WordPress successfully!")
    #         else:
    #             st.error(f"Failed to post article. Status code: {status_code}\nResponse: {response_text}")
            
    #         transalte_status_code, transalte_response_text = sendWordPrssApiTraslated(st.session_state['t_title'], st.session_state['t_text'])
    #         if transalte_status_code == 201:
    #             st.success("Translated Article posted to WordPress successfully!")
    #         else:
    #             st.error(f"Failed to post article. Status code: {transalte_status_code}\nResponse: {transalte_response_text}")

    with st.sidebar:
        st.title("Website URLs:")
        urls_input = st.text_area("Enter Website URLs (one per line)")
        urls = [url.strip() for url in urls_input.split('\n') if url.strip()]
        if st.button("Submit & Process"):
            with st.spinner("Processing..."):
                raw_text = get_website_text(urls)
                text_chunks = get_text_chunks(raw_text)
                get_vector_store(text_chunks)
                st.success("Processing Complete")

                # Export the website content to a Word document
                docx_file = export_to_docx(raw_text)
                st.success(f"Content exported to {docx_file}")
                with open(docx_file, "rb") as file:
                    st.download_button(label="Download Word Document", data=file, file_name=docx_file)

if __name__ == "__main__":
    if "logged_in" not in st.session_state:
        st.session_state.logged_in = True

    if not st.session_state.logged_in:
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            col4, col5, col6 = st.columns([1, 1, 1])
            with col5:
                st.subheader("Login")
            email = st.text_input("Enter your email:")
            password = st.text_input("Enter your password:", type="password")
            col7, col8, col9 = st.columns([4, 4, 4])
            with col8:
                if st.button("Login"):
                    if email == 'naveed' and password == '123':
                        st.session_state.logged_in = True
                        st.rerun()
                    else:
                        st.toast('Please write the email and password yourself and do not use autofill')
    else:
        main()
